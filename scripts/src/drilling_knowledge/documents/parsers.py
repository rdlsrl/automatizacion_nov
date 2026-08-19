"""Deterministic structural parsers for supported technical document formats."""

from __future__ import annotations

from dataclasses import dataclass
from html.parser import HTMLParser
import mimetypes
from pathlib import Path
import re
from decimal import Decimal
from typing import Iterable
from xml.etree import ElementTree

from drilling_knowledge.common.exceptions import ValidationError


@dataclass(frozen=True, slots=True)
class RawParagraph:
    text: str
    page_number: int | None
    heading_level: int | None = None
    source_kind: str = "paragraph"


@dataclass(frozen=True, slots=True)
class RawTable:
    rows: tuple[tuple[str, ...], ...]
    page_number: int | None
    label: str
    caption: str | None = None


@dataclass(frozen=True, slots=True)
class RawFigure:
    page_number: int | None
    label: str
    caption: str | None = None


@dataclass(frozen=True, slots=True)
class RawReference:
    text: str
    reference_type: str
    target: str | None
    page_number: int | None


@dataclass(frozen=True, slots=True)
class RawGlossaryTerm:
    term: str
    definition: str
    page_number: int | None


@dataclass(frozen=True, slots=True)
class ParsedDocumentStructure:
    parser_name: str
    title: str
    mime_type: str
    page_count: int | None
    paragraphs: tuple[RawParagraph, ...] = ()
    tables: tuple[RawTable, ...] = ()
    figures: tuple[RawFigure, ...] = ()
    references: tuple[RawReference, ...] = ()
    glossary_terms: tuple[RawGlossaryTerm, ...] = ()


class _StructuralHtmlParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.title = "Untitled HTML Document"
        self.paragraphs: list[RawParagraph] = []
        self.tables: list[RawTable] = []
        self.figures: list[RawFigure] = []
        self.references: list[RawReference] = []
        self.glossary_terms: list[RawGlossaryTerm] = []
        self._text_parts: list[str] = []
        self._current_tag: str | None = None
        self._current_href: str | None = None
        self._current_heading_level: int | None = None
        self._current_table_rows: list[list[str]] | None = None
        self._current_table_caption_parts: list[str] = []
        self._current_table_caption: str | None = None
        self._current_cell_parts: list[str] = []
        self._current_row: list[str] | None = None
        self._in_title = False
        self._in_dt = False
        self._in_caption = False
        self._pending_glossary_term: str | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_map = {key: value for key, value in attrs}
        self._current_tag = tag
        if tag == "title":
            self._in_title = True
            self._text_parts = []
        elif tag in {"p", "li"}:
            self._text_parts = []
            self._current_heading_level = None
        elif tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self._text_parts = []
            self._current_heading_level = int(tag[1])
        elif tag == "a":
            self._text_parts = []
            self._current_href = attrs_map.get("href")
        elif tag == "table":
            self._current_table_rows = []
            self._current_table_caption_parts = []
            self._current_table_caption = None
        elif tag == "tr":
            self._current_row = []
        elif tag in {"td", "th"}:
            self._current_cell_parts = []
        elif tag == "caption":
            self._text_parts = []
            self._current_table_caption_parts = []
            self._in_caption = True
        elif tag == "img":
            label = attrs_map.get("alt") or attrs_map.get("src") or "figure"
            self.figures.append(RawFigure(page_number=None, label=label.strip(), caption=attrs_map.get("title")))
        elif tag == "dt":
            self._text_parts = []
            self._in_dt = True
        elif tag == "dd":
            self._text_parts = []

    def handle_endtag(self, tag: str) -> None:
        text = " ".join(part for part in self._text_parts if part).strip()
        if tag == "title":
            if text:
                self.title = text
            self._in_title = False
        elif tag in {"p", "li"} and text:
            self.paragraphs.append(RawParagraph(text=text, page_number=None, source_kind=tag))
        elif tag in {"h1", "h2", "h3", "h4", "h5", "h6"} and text:
            self.paragraphs.append(RawParagraph(text=text, page_number=None, heading_level=self._current_heading_level, source_kind="heading"))
        elif tag == "a" and text:
            self.references.append(RawReference(text=text, reference_type="hyperlink", target=self._current_href, page_number=None))
            self._current_href = None
        elif tag == "caption":
            if self._current_table_caption_parts:
                caption_text = " ".join(part for part in self._current_table_caption_parts if part).strip()
                self._current_table_caption = caption_text or None
            self._in_caption = False
        elif tag in {"td", "th"} and self._current_row is not None:
            cell_text = " ".join(part for part in self._current_cell_parts if part).strip()
            self._current_row.append(cell_text)
            self._current_cell_parts = []
        elif tag == "tr" and self._current_table_rows is not None and self._current_row is not None:
            self._current_table_rows.append(self._current_row)
            self._current_row = None
        elif tag == "table" and self._current_table_rows:
            rows = tuple(tuple(cell.strip() for cell in row) for row in self._current_table_rows if any(cell.strip() for cell in row))
            if rows:
                self.tables.append(RawTable(rows=rows, page_number=None, label=f"table_{len(self.tables) + 1}", caption=self._current_table_caption))
            self._current_table_rows = None
            self._current_table_caption = None
        elif tag == "dt":
            self._pending_glossary_term = text or None
            self._in_dt = False
        elif tag == "dd" and text and self._pending_glossary_term:
            self.glossary_terms.append(RawGlossaryTerm(term=self._pending_glossary_term, definition=text, page_number=None))
            self._pending_glossary_term = None
        self._text_parts = []

    def handle_data(self, data: str) -> None:
        cleaned = data.strip()
        if cleaned:
            self._text_parts.append(cleaned)
            if self._current_row is not None:
                self._current_cell_parts.append(cleaned)
            if self._in_caption:
                self._current_table_caption_parts.append(cleaned)


class StructuralDocumentParser:
    def parse(self, path: Path, content_bytes: bytes) -> ParsedDocumentStructure:
        mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path, content_bytes, mime_type)
        if suffix == ".docx":
            return self._parse_docx(path, mime_type)
        if suffix == ".xlsx":
            return self._parse_xlsx(path, mime_type)
        if suffix in {".html", ".htm"}:
            return self._parse_html(content_bytes, mime_type)
        if suffix in {".md", ".markdown"}:
            return self._parse_markdown(content_bytes, mime_type)
        if suffix in {".txt", ".log"}:
            return self._parse_text(content_bytes, mime_type)
        raise ValidationError(code="unsupported_document_format", message="Unsupported document format", context={"path": str(path), "suffix": suffix})

    def _parse_pdf(self, path: Path, content_bytes: bytes, mime_type: str) -> ParsedDocumentStructure:
        try:
            from pypdf import PdfReader
        except ModuleNotFoundError as exc:
            raise ValidationError(
                code="missing_pdf_dependency",
                message="PDF parsing requires pypdf to be installed",
                context={"path": str(path), "dependency": "pypdf"},
            ) from exc
        reader = PdfReader(path)
        paragraphs: list[RawParagraph] = []
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            for paragraph in self._split_paragraphs(page_text):
                paragraphs.append(RawParagraph(text=paragraph, page_number=page_number))
        title = path.stem.replace("_", " ").strip() or path.name
        return ParsedDocumentStructure(
            parser_name="pdf",
            title=title,
            mime_type=mime_type,
            page_count=len(reader.pages),
            paragraphs=tuple(paragraphs),
        )

    def _parse_docx(self, path: Path, mime_type: str) -> ParsedDocumentStructure:
        try:
            from docx import Document as DocxDocument
        except ModuleNotFoundError as exc:
            raise ValidationError(
                code="missing_docx_dependency",
                message="DOCX parsing requires python-docx to be installed",
                context={"path": str(path), "dependency": "python-docx"},
            ) from exc
        doc = DocxDocument(path)
        paragraphs: list[RawParagraph] = []
        tables: list[RawTable] = []
        references: list[RawReference] = []
        title = path.stem.replace("_", " ").strip() or path.name
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            heading_level = None
            if style_name.startswith("heading"):
                match = re.search(r"(\d+)", style_name)
                heading_level = int(match.group(1)) if match else 1
            paragraphs.append(RawParagraph(text=text, page_number=None, heading_level=heading_level, source_kind="docx_paragraph"))
            if paragraph.text.lower().startswith("see "):
                references.append(RawReference(text=paragraph.text, reference_type="inline_reference", target=None, page_number=None))
        for index, table in enumerate(doc.tables, start=1):
            rows = tuple(tuple(cell.text.strip() for cell in row.cells) for row in table.rows)
            if any(any(cell for cell in row) for row in rows):
                tables.append(RawTable(rows=rows, page_number=None, label=f"table_{index}"))
        if doc.core_properties.title:
            title = doc.core_properties.title.strip() or title
        return ParsedDocumentStructure(
            parser_name="docx",
            title=title,
            mime_type=mime_type,
            page_count=None,
            paragraphs=tuple(paragraphs),
            tables=tuple(tables),
            references=tuple(references),
        )

    def _parse_xlsx(self, path: Path, mime_type: str) -> ParsedDocumentStructure:
        try:
            from openpyxl import load_workbook
        except ModuleNotFoundError as exc:
            raise ValidationError(
                code="missing_xlsx_dependency",
                message="XLSX parsing requires openpyxl to be installed",
                context={"path": str(path), "dependency": "openpyxl"},
            ) from exc
        workbook = load_workbook(path, read_only=True, data_only=True)
        paragraphs: list[RawParagraph] = []
        tables: list[RawTable] = []
        title = path.stem.replace("_", " ").strip() or path.name
        for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
            sheet_name = (sheet.title or f"Sheet{sheet_index}").strip() or f"Sheet{sheet_index}"
            paragraphs.append(RawParagraph(text=sheet_name, page_number=sheet_index, heading_level=1, source_kind="sheet_heading"))
            rows = self._xlsx_rows(sheet)
            if not rows:
                continue
            header = rows[0]
            column_count = max(len(row) for row in rows)
            data_row_count = max(len(rows) - 1, 0)
            paragraphs.append(
                RawParagraph(
                    text=(
                        f"Sheet {sheet_name} from file {path.name}. "
                        f"Columns: {' | '.join(header)}. Data rows: {data_row_count}."
                    ),
                    page_number=sheet_index,
                    source_kind="table_caption",
                )
            )
            tables.append(
                RawTable(
                    rows=rows,
                    page_number=sheet_index,
                    label=f"sheet:{sheet_name}",
                    caption=f"file={path.name}; sheet={sheet_name}; columns={column_count}; rows={len(rows)}",
                )
            )
        workbook.close()
        return ParsedDocumentStructure(
            parser_name="xlsx",
            title=title,
            mime_type=mime_type,
            page_count=len(workbook.worksheets),
            paragraphs=tuple(paragraphs),
            tables=tuple(tables),
        )

    def _parse_html(self, content_bytes: bytes, mime_type: str) -> ParsedDocumentStructure:
        parser = _StructuralHtmlParser()
        parser.feed(content_bytes.decode("utf-8"))
        return ParsedDocumentStructure(
            parser_name="html",
            title=parser.title,
            mime_type=mime_type,
            page_count=None,
            paragraphs=tuple(parser.paragraphs),
            tables=tuple(parser.tables),
            figures=tuple(parser.figures),
            references=tuple(parser.references),
            glossary_terms=tuple(parser.glossary_terms),
        )

    def _parse_markdown(self, content_bytes: bytes, mime_type: str) -> ParsedDocumentStructure:
        text = content_bytes.decode("utf-8")
        lines = text.splitlines()
        paragraphs: list[RawParagraph] = []
        tables: list[RawTable] = []
        figures: list[RawFigure] = []
        references: list[RawReference] = []
        glossary_terms: list[RawGlossaryTerm] = []
        buffer: list[str] = []
        title = "Untitled Markdown Document"
        index = 0
        while index < len(lines):
            line = lines[index].rstrip()
            stripped = line.strip()
            if stripped.startswith("#"):
                if buffer:
                    paragraphs.extend(self._paragraphs_from_buffer(buffer))
                    buffer = []
                level = len(stripped) - len(stripped.lstrip("#"))
                heading_text = stripped[level:].strip()
                if level == 1 and title == "Untitled Markdown Document":
                    title = heading_text
                paragraphs.append(RawParagraph(text=heading_text, page_number=None, heading_level=level, source_kind="heading"))
                index += 1
                continue
            if stripped.startswith("!["):
                alt_text = stripped[2:].split("]", 1)[0]
                figures.append(RawFigure(page_number=None, label=alt_text or f"figure_{len(figures) + 1}"))
                index += 1
                continue
            inline_links = list(re.finditer(r"\[(.+?)\]\((.+?)\)", stripped))
            for link_match in inline_links:
                references.append(
                    RawReference(
                        text=link_match.group(1),
                        reference_type="hyperlink",
                        target=link_match.group(2),
                        page_number=None,
                    )
                )
            if inline_links and stripped == inline_links[0].group(0):
                index += 1
                continue
            if stripped.lower().startswith("glossary:"):
                if buffer:
                    paragraphs.extend(self._paragraphs_from_buffer(buffer))
                    buffer = []
                index += 1
                while index < len(lines) and lines[index].strip():
                    term_line = lines[index].strip()
                    if ":" in term_line:
                        term, definition = term_line.split(":", 1)
                        glossary_terms.append(RawGlossaryTerm(term=term.strip(), definition=definition.strip(), page_number=None))
                    index += 1
                continue
            if stripped.startswith("|"):
                if buffer:
                    paragraphs.extend(self._paragraphs_from_buffer(buffer))
                    buffer = []
                table_lines = []
                while index < len(lines) and lines[index].strip().startswith("|"):
                    table_lines.append(lines[index].strip())
                    index += 1
                rows = []
                for row_index, table_line in enumerate(table_lines):
                    cells = tuple(cell.strip() for cell in table_line.strip("|").split("|"))
                    if row_index == 1 and all(set(cell) <= {"-", ":"} for cell in cells):
                        continue
                    rows.append(cells)
                if rows:
                    tables.append(RawTable(rows=tuple(rows), page_number=None, label=f"table_{len(tables) + 1}"))
                continue
            if not stripped:
                if buffer:
                    paragraphs.extend(self._paragraphs_from_buffer(buffer))
                    buffer = []
                index += 1
                continue
            buffer.append(stripped)
            index += 1
        if buffer:
            paragraphs.extend(self._paragraphs_from_buffer(buffer))
        return ParsedDocumentStructure(
            parser_name="markdown",
            title=title,
            mime_type=mime_type,
            page_count=None,
            paragraphs=tuple(paragraphs),
            tables=tuple(tables),
            figures=tuple(figures),
            references=tuple(references),
            glossary_terms=tuple(glossary_terms),
        )

    def _parse_text(self, content_bytes: bytes, mime_type: str) -> ParsedDocumentStructure:
        text = content_bytes.decode("utf-8")
        raw_paragraphs = self._split_paragraphs(text)
        title = raw_paragraphs[0] if raw_paragraphs else "Untitled Text Document"
        body_paragraphs = raw_paragraphs[1:] if len(raw_paragraphs) > 1 else raw_paragraphs
        paragraphs = tuple(RawParagraph(text=paragraph, page_number=None) for paragraph in body_paragraphs)
        return ParsedDocumentStructure(
            parser_name="text",
            title=title,
            mime_type=mime_type,
            page_count=None,
            paragraphs=paragraphs,
        )

    def _split_paragraphs(self, text: str) -> list[str]:
        return [paragraph.strip() for paragraph in re.split(r"\n\s*\n", text) if paragraph.strip()]

    def _paragraphs_from_buffer(self, buffer: Iterable[str]) -> list[RawParagraph]:
        text = " ".join(line.strip() for line in buffer if line.strip())
        return [RawParagraph(text=text, page_number=None)] if text else []

    def _xlsx_rows(self, sheet) -> tuple[tuple[str, ...], ...]:
        raw_rows: list[tuple[str, ...]] = []
        max_width = 0
        for row in sheet.iter_rows(values_only=True):
            values = tuple(self._stringify_spreadsheet_cell(cell) for cell in row)
            if any(value != "" for value in values):
                raw_rows.append(values)
                max_width = max(max_width, len(values))
        if not raw_rows or max_width == 0:
            return ()
        normalized_rows: list[tuple[str, ...]] = []
        for row in raw_rows:
            padded = row + ("",) * (max_width - len(row))
            normalized_rows.append(padded)
        return tuple(normalized_rows)

    @staticmethod
    def _stringify_spreadsheet_cell(value: object) -> str:
        if value is None:
            return ""
        if isinstance(value, bool):
            return "TRUE" if value else "FALSE"
        if isinstance(value, float):
            return format(value, ".15g")
        if isinstance(value, Decimal):
            return format(value, "f")
        return str(value).strip()