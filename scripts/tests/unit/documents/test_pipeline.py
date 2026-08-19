from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from drilling_knowledge.documents import DocumentKnowledgeAcquisitionEngine, DocumentMetadata


class DocumentKnowledgeAcquisitionEngineTests(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = DocumentKnowledgeAcquisitionEngine.create()
        self.metadata = DocumentMetadata(
            author="Test Author",
            manufacturer="Generic OEM",
            model="Generic Model",
            version_label="1.0",
            language="en",
            authority_level="reference",
            document_type="manual",
            source="test_fixture",
            license_name="internal",
        )

    def test_ingests_markdown_with_sections_tables_references_and_glossary(self) -> None:
        markdown = """# OEM Manual\n\nOverview paragraph.\n\n## 5.1 Connections\n\nSee [WITS](https://example.com/wits). Refer to Figure 1, Table 1 and Section 5.1.\n\n![Figure 1](figure-1.png)\n\n| Name | Value |\n| --- | --- |\n| Pump | Online |\n\nGlossary:\nSPP: Standpipe pressure\n"""
        with temporary_file("manual.md", markdown.encode("utf-8")) as path:
            result = self.engine.ingest(path, self.metadata)

        self.assertEqual(result.snapshot.document.title, "OEM Manual")
        self.assertGreaterEqual(len(result.snapshot.sections), 2)
        self.assertEqual(len(result.snapshot.tables), 1)
        self.assertEqual(len(result.snapshot.figures), 1)
        self.assertEqual(len(result.snapshot.references), 4)
        self.assertEqual(len(result.snapshot.glossary_terms), 1)
        self.assertTrue(all(fragment.trace.document_id == result.snapshot.document.entity_id for fragment in result.snapshot.fragments))
        self.assertTrue(all(fragment.content_hash for fragment in result.snapshot.fragments))

        reference_map = {reference.reference_text: reference for reference in result.snapshot.references}
        target_section = next(section for section in result.snapshot.sections if section.title == "5.1 Connections")
        self.assertEqual(reference_map["Figure 1"].resolved_figure_id, result.snapshot.figures[0].entity_id)
        self.assertEqual(reference_map["Table 1"].resolved_table_id, result.snapshot.tables[0].entity_id)
        self.assertEqual(reference_map["Section 5.1"].resolved_section_id, target_section.entity_id)

    def test_ingests_html_with_tables_figures_and_glossary(self) -> None:
        html = b"""<html><head><title>Instrument Manual</title></head><body><h1>Instrument Manual</h1><p>Body paragraph.</p><a href='#figure1'>Figure ref</a><img src='fig1.png' alt='Figure 1'/><table><tr><th>Name</th><th>Value</th></tr><tr><td>Flow</td><td>On</td></tr></table><dl><dt>PLC</dt><dd>Programmable logic controller</dd></dl></body></html>"""
        with temporary_file("instrument.html", html) as path:
            result = self.engine.ingest(path, self.metadata)

        self.assertEqual(result.snapshot.document.title, "Instrument Manual")
        self.assertEqual(len(result.snapshot.tables), 1)
        self.assertEqual(len(result.snapshot.figures), 1)
        self.assertEqual(len(result.snapshot.references), 1)
        self.assertEqual(len(result.snapshot.glossary_terms), 1)

    def test_ingests_plain_text(self) -> None:
        text = b"Plain text title\n\nFirst paragraph.\n\nSecond paragraph."
        with temporary_file("notes.txt", text) as path:
            result = self.engine.ingest(path, self.metadata)

        self.assertEqual(result.snapshot.version.parser_name, "text")
        self.assertEqual(len(result.snapshot.fragments), 2)

    def test_ingests_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "manual.docx"
            document = DocxDocument()
            document.core_properties.title = "DOCX Manual"
            document.add_heading("DOCX Manual", level=1)
            document.add_paragraph("Document body paragraph.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Pump"
            table.cell(1, 1).text = "Ready"
            document.save(path)

            result = self.engine.ingest(path, self.metadata)

        self.assertEqual(result.snapshot.document.title, "DOCX Manual")
        self.assertEqual(result.snapshot.version.parser_name, "docx")
        self.assertEqual(len(result.snapshot.tables), 1)
        self.assertGreaterEqual(len(result.snapshot.sections), 2)

    def test_ingests_pdf(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "manual.pdf"
            pdf = canvas.Canvas(str(path))
            pdf.drawString(72, 720, "PDF Manual")
            pdf.drawString(72, 700, "First paragraph in pdf.")
            pdf.drawString(72, 680, "Second paragraph in pdf.")
            pdf.save()

            result = self.engine.ingest(path, self.metadata)

        self.assertEqual(result.snapshot.version.parser_name, "pdf")
        self.assertEqual(result.snapshot.version.page_count, 1)
        self.assertGreaterEqual(len(result.snapshot.fragments), 1)

    def test_ingests_xlsx_preserving_sheet_headers_and_rows(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "dataset.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "DPR Data"
            sheet.append(["Depth ft", "Pressure psi", "Flow gpm"])
            sheet.append([1000, 2500, 120])
            sheet.append([1005, 2525, 118])
            workbook.save(path)

            result = self.engine.ingest(path, self.metadata)

        self.assertEqual(result.snapshot.document.title, "dataset")
        self.assertEqual(result.snapshot.version.parser_name, "xlsx")
        self.assertEqual(result.snapshot.version.page_count, 1)
        self.assertEqual(len(result.snapshot.tables), 1)
        self.assertIn("sheet:DPR Data", result.snapshot.tables[0].label)
        self.assertEqual(result.snapshot.tables[0].rows[0], ("Depth ft", "Pressure psi", "Flow gpm"))
        self.assertEqual(result.snapshot.tables[0].rows[1], ("1000", "2500", "120"))
        self.assertTrue(any(fragment.fragment_type == "table" for fragment in result.snapshot.fragments))

    def test_persistence_is_stable_for_same_document_version(self) -> None:
        markdown = b"# Stable Document\n\nBody paragraph."
        with temporary_file("stable.md", markdown) as path:
            first = self.engine.ingest(path, self.metadata)
            second = self.engine.ingest(path, self.metadata)

        self.assertEqual(first.snapshot.version.entity_id, second.snapshot.version.entity_id)
        self.assertEqual(len(second.repository.versions), 1)
        self.assertEqual(len(second.repository.fragments), len(first.snapshot.fragments))
        self.assertEqual(
            [fragment.content_hash for fragment in first.snapshot.fragments],
            [fragment.content_hash for fragment in second.snapshot.fragments],
        )


class TemporaryFile:
    def __init__(self, name: str, payload: bytes) -> None:
        self.name = name
        self.payload = payload
        self.temp_dir = tempfile.TemporaryDirectory()
        self.path = Path(self.temp_dir.name) / name

    def __enter__(self) -> Path:
        self.path.write_bytes(self.payload)
        return self.path

    def __exit__(self, exc_type, exc, tb) -> None:
        self.temp_dir.cleanup()


def temporary_file(name: str, payload: bytes) -> TemporaryFile:
    return TemporaryFile(name, payload)